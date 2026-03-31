import streamlit as st
import google.generativeai as genai
import os
import json
import re
import hashlib
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import requests
from supabase import create_client, Client

# ========== KONFIGURASI ==========
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

st.set_page_config(
    page_title="Generator Soal Multi Mapel",
    page_icon="📖",
    layout="wide"
)

# ========== SUPABASE KONEKSI ==========
# GANTI DENGAN DATA DARI SUPABASE ANDA!
SUPABASE_URL = "https://mzixdthfukblrhnifunl.supabase.co"  # ganti dengan URL Anda
SUPABASE_KEY = "sb_publishable_0yIRih0UtYcl47kPCaleCQ_bMZmEaFm"  # ganti dengan anon key Anda

@st.cache_resource
def init_supabase():
    return create_client(SUPABASE_URL, SUPABASE_KEY)

supabase = init_supabase()

# ========== FUNGSI AUTHENTIKASI ==========
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def register_user(username, password):
    try:
        hashed_pw = hash_password(password)
        supabase.table("users").insert({
            "username": username,
            "password": hashed_pw
        }).execute()
        return True, "Registrasi berhasil! Silakan login."
    except Exception as e:
        if "duplicate key" in str(e).lower():
            return False, "Username sudah digunakan!"
        return False, f"Gagal registrasi: {e}"

def login_user(username, password):
    try:
        hashed_pw = hash_password(password)
        response = supabase.table("users").select("*").eq("username", username).eq("password", hashed_pw).execute()
        if response.data:
            return True, response.data[0]["id"]
        return False, "Username atau password salah!"
    except Exception as e:
        return False, f"Gagal login: {e}"

# ========== FUNGSI DATABASE MAPEL ==========
def load_mapel(user_id):
    try:
        response = supabase.table("mapel").select("*").eq("user_id", user_id).order("id").execute()
        return response.data
    except Exception as e:
        st.error(f"Gagal load mapel: {e}")
        return []

def create_mapel(nama_mapel, user_id):
    try:
        supabase.table("mapel").insert({
            "nama_mapel": nama_mapel,
            "user_id": user_id
        }).execute()
        return True, "Mapel berhasil dibuat!"
    except Exception as e:
        if "duplicate key" in str(e).lower():
            return False, "Anda sudah memiliki mapel dengan nama ini!"
        return False, f"Gagal membuat mapel: {e}"

def delete_mapel(mapel_id, user_id):
    try:
        supabase.table("mapel").delete().eq("id", mapel_id).eq("user_id", user_id).execute()
        return True, "Mapel berhasil dihapus!"
    except Exception as e:
        return False, f"Gagal menghapus mapel: {e}"

# ========== FUNGSI DATABASE KI/KD ==========
def load_ki_kd(mapel_id):
    try:
        response = supabase.table("ki_kd").select("*").eq("mapel_id", mapel_id).order("kelas", desc=False).order("id").execute()
        return response.data
    except Exception as e:
        return []

def save_ki_kd(mapel_id, kelas, ki, kd, keterangan):
    try:
        supabase.table("ki_kd").insert({
            "mapel_id": mapel_id,
            "kelas": kelas,
            "ki": ki,
            "kd": kd,
            "keterangan": keterangan
        }).execute()
        return True
    except Exception as e:
        st.error(f"Gagal menyimpan KI/KD: {e}")
        return False

def delete_ki_kd(ki_kd_id, mapel_id):
    try:
        supabase.table("ki_kd").delete().eq("id", ki_kd_id).eq("mapel_id", mapel_id).execute()
        return True
    except Exception as e:
        st.error(f"Gagal menghapus KI/KD: {e}")
        return False

def save_hasil_soal(mapel_id, kelas, topik, soal_data):
    try:
        supabase.table("hasil_soal").insert({
            "mapel_id": mapel_id,
            "kelas": kelas,
            "topik": topik,
            "soal_data": soal_data
        }).execute()
    except Exception as e:
        pass

# ========== INISIALISASI SESSION STATE ==========
def init_session_state():
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False
    if 'user_id' not in st.session_state:
        st.session_state.user_id = None
    if 'username' not in st.session_state:
        st.session_state.username = None
    if 'mapel_list' not in st.session_state:
        st.session_state.mapel_list = []
    if 'selected_mapel_id' not in st.session_state:
        st.session_state.selected_mapel_id = None
    if 'selected_mapel_nama' not in st.session_state:
        st.session_state.selected_mapel_nama = None
    if 'ki_kd_list' not in st.session_state:
        st.session_state.ki_kd_list = []
    if 'show_generator' not in st.session_state:
        st.session_state.show_generator = False
    if 'hasil_soal' not in st.session_state:
        st.session_state.hasil_soal = None

# ========== FUNGSI GENERATE SOAL (disingkat karena panjang) ==========
def buat_soal(api_key, prompt, ai_provider, deepseek_model="", openai_model="", maia_endpoint="", maia_model=""):
    try:
        if ai_provider == "Gemini (Google)":
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-2.0-flash')
            response = model.generate_content(prompt)
            response_text = response.text
        elif ai_provider == "DeepSeek":
            if not OPENAI_AVAILABLE:
                return None, "Library openai belum terinstall"
            client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com/v1")
            response = client.chat.completions.create(
                model=deepseek_model,
                messages=[{"role": "system", "content": "Anda adalah guru ahli. Buat soal dalam format JSON."}, {"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=4000
            )
            response_text = response.choices[0].message.content
        elif ai_provider == "OpenAI GPT":
            if not OPENAI_AVAILABLE:
                return None, "Library openai belum terinstall"
            client = OpenAI(api_key=api_key, base_url="https://api.openai.com/v1")
            response = client.chat.completions.create(
                model=openai_model,
                messages=[{"role": "system", "content": "Anda adalah guru ahli. Buat soal dalam format JSON."}, {"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=4000
            )
            response_text = response.choices[0].message.content
        elif ai_provider == "Maia Router":
            if not maia_endpoint:
                maia_endpoint = "https://api.maiarouter.com/v1/chat/completions"
            headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
            payload = {"model": maia_model, "messages": [{"role": "system", "content": "Anda adalah guru ahli. Buat soal dalam format JSON."}, {"role": "user", "content": prompt}], "temperature": 0.7, "max_tokens": 4000}
            response = requests.post(maia_endpoint, headers=headers, json=payload, timeout=60)
            if response.status_code == 200:
                response_data = response.json()
                response_text = response_data['choices'][0]['message']['content']
            else:
                return None, f"API error: {response.status_code}"
        else:
            return None, f"Provider tidak dikenal"
        
        match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if match:
            return json.loads(match.group()), None
        return None, "Tidak ada format JSON"
    except Exception as e:
        return None, str(e)

def build_prompt(ki, kd, kelas, topik, bloom_levels, tipe_config, sumber_materi, teks_bacaan, terjemahan, mode_pembuatan, contoh_soal, tampilan_bacaan, total_soal):
    materi_khusus = ""
    if sumber_materi == "Dengan Materi Khusus (Teks/Topik)" and teks_bacaan:
        instruksi = "Teks bacaan WAJIB ditampilkan di AWAL." if tampilan_bacaan == "Teks bacaan ditampilkan di atas semua soal" else "Teks bacaan HANYA untuk soal berbasis teks." if tampilan_bacaan == "Teks bacaan hanya ditampilkan pada soal berbasis teks saja" else "Teks bacaan TIDAK ditampilkan."
        materi_khusus = f"Teks: {teks_bacaan}\nTerjemahan: {terjemahan}\n{instruksi}"
    
    bagian_contoh = ""
    if mode_pembuatan == "Ikuti Contoh Soal (Parafrase)" and contoh_soal:
        bagian_contoh = f"IKUTI FORMAT CONTOH:\n{contoh_soal}\nBuat {total_soal} soal dengan format SAMA."
    
    prompt = f"""Anda guru ahli SMP/MTs.
KI: {ki}
KD: {kd}
Kelas: {kelas}
Topik: {topik}
{materi_khusus}
{bagian_contoh}
Level Bloom: {', '.join(bloom_levels)}
Tipe Soal: PG Biasa:{tipe_config.get('jumlah_pg',0)}, PG Kompleks:{tipe_config.get('jumlah_pg_kompleks',0)}, Benar/Salah:{tipe_config.get('jumlah_benar_salah',0)}, PG Teks:{tipe_config.get('jumlah_teks',0)}
Output JSON: {{"metadata":{{"kelas":"{kelas}","topik":"{topik}"}},"soal":[{{"nomor":1,"tipe":"Pilihan Ganda Biasa","bloom_level":"C1","teks_arab":"","teks_indonesia":"","pilihan":{{"A":"","B":"","C":"","D":""}},"jawaban":"A","pembahasan":""}}]}}"""
    return prompt

def display_question(soal, nomor, tampilan_bacaan, teks_bacaan_global, terjemahan_global, sudah_tampil):
    if tampilan_bacaan == "Teks bacaan ditampilkan di atas semua soal" and nomor == 1 and not sudah_tampil and teks_bacaan_global:
        st.markdown("### 📖 Teks Bacaan")
        st.markdown(f"<div dir='rtl' style='background:#e9ecef;padding:15px;border-radius:10px'>{teks_bacaan_global}</div>", unsafe_allow_html=True)
        st.markdown("---")
        return True
    st.markdown(f"**{nomor}. {soal.get('teks_indonesia', soal)}**")
    st.divider()
    return False

def export_to_word(soal_data, ki, kd, kelas, topik, sumber_materi, teks_bacaan=None):
    doc = Document()
    doc.add_heading('SOAL', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Kelas: {kelas}")
    doc.add_paragraph(f"Topik: {topik}")
    doc.add_paragraph()
    if ki:
        doc.add_heading('KI', level=2)
        doc.add_paragraph(ki)
    if kd:
        doc.add_heading('KD', level=2)
        doc.add_paragraph(kd)
    doc.add_heading('SOAL', level=1)
    for i, soal in enumerate(soal_data.get('soal', []), 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(soal.get('teks_indonesia', ''))
    return doc

def buat_kunci_jawaban_word(soal_data):
    doc = Document()
    doc.add_heading('KUNCI JAWABAN', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'No'
    table.rows[0].cells[1].text = 'Jawaban'
    for soal in soal_data.get('soal', []):
        row = table.add_row().cells
        row[0].text = str(soal.get('nomor', ''))
        jawaban = soal.get('jawaban', '')
        if isinstance(jawaban, list):
            jawaban = ', '.join(jawaban)
        row[1].text = str(jawaban)
    return doc

# ========== HALAMAN LOGIN (DIPERBAIKI) ==========
def show_login_page():
    # CSS untuk styling
    st.markdown("""
        <style>
        /* Mengatur container login agar di tengah */
        .login-container {
            max-width: 400px;
            margin: 0 auto;
            padding: 2rem;
            background-color: white;
            border-radius: 12px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
        }
        /* Mengatur judul */
        .login-title {
            text-align: center;
            margin-bottom: 2rem;
        }
        /* Mengatur tombol */
        .stButton button {
            width: 100%;
            border-radius: 8px;
            font-weight: 500;
        }
        /* Mengatur input field */
        .stTextInput input {
            border-radius: 8px;
        }
        </style>
    """, unsafe_allow_html=True)
    
    # Container tengah
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown('<div class="login-container">', unsafe_allow_html=True)
        
        # Logo/Icon
        st.markdown('<div class="login-title">', unsafe_allow_html=True)
        st.markdown("# 📖 Generator Soal")
        st.markdown("##### Multi Mapel | Multi AI")
        st.markdown('</div>', unsafe_allow_html=True)
        
        tab1, tab2 = st.tabs(["🔐 Login", "📝 Registrasi"])
        
        with tab1:
            with st.form("login_form", clear_on_submit=False):
                username = st.text_input("Username", placeholder="Masukkan username Anda", key="login_username")
                password = st.text_input("Password", type="password", placeholder="Masukkan password", key="login_password")
                
                st.markdown("<br>", unsafe_allow_html=True)
                
                submitted = st.form_submit_button("Login", use_container_width=True, type="primary")
                
                if submitted:
                    if username and password:
                        success, result = login_user(username, password)
                        if success:
                            st.session_state.logged_in = True
                            st.session_state.user_id = result
                            st.session_state.username = username
                            st.session_state.mapel_list = load_mapel(result)
                            st.success(f"Selamat datang, {username}!")
                            st.rerun()
                        else:
                            st.error(result)
                    else:
                        st.error("Masukkan username dan password!")
        
        with tab2:
            with st.form("register_form", clear_on_submit=False):
                new_username = st.text_input("Username Baru", placeholder="Pilih username", key="reg_username")
                new_password = st.text_input("Password Baru", type="password", placeholder="Pilih password", key="reg_password")
                confirm_password = st.text_input("Konfirmasi Password", type="password", placeholder="Ulangi password", key="reg_confirm")
                
                st.markdown("<br>", unsafe_allow_html=True)
                
                submitted = st.form_submit_button("Registrasi", use_container_width=True)
                
                if submitted:
                    if not new_username or not new_password:
                        st.error("Username dan password harus diisi!")
                    elif new_password != confirm_password:
                        st.error("Password dan konfirmasi tidak sama!")
                    else:
                        success, msg = register_user(new_username, new_password)
                        if success:
                            st.success(msg)
                        else:
                            st.error(msg)
        
        # Footer
        st.markdown("---")
        st.caption("✨ Ditenagai Gemini | DeepSeek | OpenAI | Maia Router")
        st.caption("🔒 Data Anda aman dan tersimpan di database")
        
        st.markdown('</div>', unsafe_allow_html=True)

# ========== DASHBOARD MAPEL ==========
def show_dashboard():
    st.title("📚 Dashboard Mapel")
    st.markdown(f"Selamat datang, **{st.session_state.username}**!")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("📁 Mapel Saya")
        if not st.session_state.mapel_list:
            st.info("Belum ada mapel. Buat mapel baru!")
        else:
            for mapel in st.session_state.mapel_list:
                col_a, col_b, col_c = st.columns([3, 1, 1])
                with col_a:
                    st.markdown(f"**📖 {mapel['nama_mapel']}**")
                with col_b:
                    if st.button("Pilih", key=f"select_{mapel['id']}"):
                        st.session_state.selected_mapel_id = mapel['id']
                        st.session_state.selected_mapel_nama = mapel['nama_mapel']
                        st.session_state.ki_kd_list = load_ki_kd(mapel['id'])
                        st.session_state.show_generator = False
                        st.rerun()
                with col_c:
                    if st.button("🗑️", key=f"del_{mapel['id']}"):
                        success, msg = delete_mapel(mapel['id'], st.session_state.user_id)
                        if success:
                            st.session_state.mapel_list = load_mapel(st.session_state.user_id)
                            st.success(msg)
                            st.rerun()
                        else:
                            st.error(msg)
                st.divider()
    
    with col2:
        st.subheader("➕ Buat Mapel Baru")
        with st.form("new_mapel_form"):
            nama_mapel_baru = st.text_input("Nama Mapel", placeholder="Contoh: Bahasa Arab")
            if st.form_submit_button("Buat Mapel", use_container_width=True):
                if nama_mapel_baru:
                    success, msg = create_mapel(nama_mapel_baru, st.session_state.user_id)
                    if success:
                        st.session_state.mapel_list = load_mapel(st.session_state.user_id)
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)
    
    if st.button("🚪 Logout", use_container_width=True):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

# ========== HALAMAN KELOLA KI/KD ==========
def show_manage_ki_kd():
    st.title(f"📖 Kelola KI/KD - {st.session_state.selected_mapel_nama}")
    
    if st.button("← Kembali ke Dashboard"):
        st.session_state.selected_mapel_id = None
        st.session_state.selected_mapel_nama = None
        st.session_state.ki_kd_list = []
        st.rerun()
    
    with st.expander("➕ Tambah KI/KD Baru", expanded=True):
        kelas_baru = st.selectbox("Kelas", ["7", "8", "9"])
        keterangan_baru = st.text_input("Keterangan", placeholder="Bab 1")
        ki_baru = st.text_area("KI", height=80)
        kd_baru = st.text_area("KD", height=80)
        
        if st.button("💾 Simpan KI/KD", use_container_width=True):
            if ki_baru and kd_baru:
                if save_ki_kd(st.session_state.selected_mapel_id, kelas_baru, ki_baru, kd_baru, keterangan_baru):
                    st.session_state.ki_kd_list = load_ki_kd(st.session_state.selected_mapel_id)
                    st.success("✅ Tersimpan!")
                    st.rerun()
                else:
                    st.error("Gagal menyimpan!")
            else:
                st.error("KI dan KD harus diisi!")
    
    st.subheader("📋 Daftar KI/KD")
    if not st.session_state.ki_kd_list:
        st.info("Belum ada KI/KD.")
    else:
        for item in st.session_state.ki_kd_list:
            col_btn, col_info = st.columns([1, 5])
            with col_info:
                st.markdown(f"**Kelas {item['kelas']}** {f'- {item.get("keterangan", "")}' if item.get("keterangan") else ''}")
                st.caption(f"KI: {item['ki'][:80]}...")
                st.caption(f"KD: {item['kd'][:80]}...")
            with col_btn:
                if st.button("🗑️", key=f"del_kd_{item['id']}"):
                    if delete_ki_kd(item['id'], st.session_state.selected_mapel_id):
                        st.session_state.ki_kd_list = load_ki_kd(st.session_state.selected_mapel_id)
                        st.rerun()
            st.divider()
    
    if st.session_state.ki_kd_list:
        if st.button("🚀 Lanjut ke Generator Soal", use_container_width=True, type="primary"):
            st.session_state.show_generator = True
            st.rerun()

# ========== GENERATOR SOAL ==========
def show_generator():
    st.title(f"📝 Generator Soal - {st.session_state.selected_mapel_nama}")
    
    if st.button("← Kembali ke KI/KD"):
        st.session_state.show_generator = False
        st.rerun()
    
    # Pilih KI/KD
    ki_kd_options = []
    for item in st.session_state.ki_kd_list:
        label = f"Kelas {item['kelas']}: {item.get('keterangan', item['kd'][:50])}..."
        ki_kd_options.append({"label": label, "kelas": item["kelas"], "ki": item["ki"], "kd": item["kd"]})
    
    if not ki_kd_options:
        st.warning("Belum ada KI/KD. Silakan tambahkan dulu!")
        return
    
    selected_index = st.selectbox("Pilih KI/KD", range(len(ki_kd_options)), format_func=lambda x: ki_kd_options[x]["label"])
    selected = ki_kd_options[selected_index]
    kelas = selected["kelas"]
    ki = selected["ki"]
    kd = selected["kd"]
    
    st.info(f"**KI:** {ki[:150]}...")
    st.info(f"**KD:** {kd[:150]}...")
    
    # Sidebar settings
    with st.sidebar:
        st.header("⚙️ Pengaturan")
        
        ai_provider = st.selectbox("AI Provider", ["Gemini (Google)", "DeepSeek", "OpenAI GPT", "Maia Router"])
        
        if ai_provider == "Gemini (Google)":
            api_key = st.text_input("API Key Gemini", type="password")
            deepseek_model = openai_model = maia_endpoint = maia_model = ""
        elif ai_provider == "DeepSeek":
            api_key = st.text_input("API Key DeepSeek", type="password")
            deepseek_model = st.selectbox("Model", ["deepseek-chat", "deepseek-reasoner"])
            openai_model = maia_endpoint = maia_model = ""
        elif ai_provider == "OpenAI GPT":
            api_key = st.text_input("API Key OpenAI", type="password")
            openai_model = st.selectbox("Model", ["gpt-4o", "gpt-4o-mini"])
            deepseek_model = maia_endpoint = maia_model = ""
        else:
            api_key = st.text_input("API Key Maia Router", type="password")
            maia_endpoint = st.text_input("Endpoint", placeholder="https://api.maiarouter.com/v1/chat/completions")
            maia_model = st.selectbox("Model", ["gpt-4", "claude-3", "auto"])
            deepseek_model = openai_model = ""
        
        mode_pembuatan = st.radio("Mode", ["AI Bebas", "Ikuti Contoh Soal"])
        contoh_soal = ""
        if mode_pembuatan == "Ikuti Contoh Soal":
            contoh_soal = st.text_area("Contoh Soal", height=150)
        
        sumber_materi = st.radio("Sumber Materi", ["Sesuai KI/KD", "Dengan Materi Khusus"])
        teks_bacaan = terjemahan = topik_khusus = ""
        tampilan_bacaan = "Tidak ditampilkan"
        if sumber_materi == "Dengan Materi Khusus":
            teks_bacaan = st.text_area("Teks Bacaan", height=100)
            terjemahan = st.text_area("Terjemahan", height=60)
            topik_khusus = st.text_input("Topik Khusus")
            tampilan_bacaan = st.radio("Tampilan", ["Di atas semua soal", "Hanya soal berbasis teks", "Tidak ditampilkan"])
        
        topik = topik_khusus if topik_khusus else st.text_input("Topik Umum")
        
        tipe_pg = st.checkbox("PG Biasa", True)
        tipe_pg_kompleks = st.checkbox("PG Kompleks", True)
        tipe_benar_salah = st.checkbox("Benar/Salah", True)
        tipe_teks = st.checkbox("PG Berbasis Teks", True)
        
        bloom_levels = st.multiselect("Bloom", ["C1 - Mengingat", "C2 - Memahami"], default=["C1 - Mengingat"])
        
        total_soal = st.number_input("Jumlah Soal", 1, 20, 5)
        jumlah_pg = total_soal if tipe_pg else 0
        jumlah_pg_kompleks = total_soal if tipe_pg_kompleks else 0
        jumlah_benar_salah = total_soal if tipe_benar_salah else 0
        jumlah_teks = total_soal if tipe_teks else 0
        
        enable_edit = st.checkbox("Tampilkan editor", True)
        
        tombol = st.button("🚀 Buat Soal!", type="primary", use_container_width=True)
    
    # Proses generate soal
    if tombol:
        if not api_key:
            st.error("Masukkan API Key!")
        elif not ki or not kd:
            st.error("Pilih KI/KD!")
        else:
            tipe_config = {
                'pg_biasa': tipe_pg, 'pg_kompleks': tipe_pg_kompleks,
                'benar_salah': tipe_benar_salah, 'pg_teks': tipe_teks,
                'jumlah_pg': jumlah_pg, 'jumlah_pg_kompleks': jumlah_pg_kompleks,
                'jumlah_benar_salah': jumlah_benar_salah, 'jumlah_teks': jumlah_teks
            }
            
            prompt = build_prompt(
                ki, kd, kelas, topik, bloom_levels, tipe_config,
                sumber_materi, teks_bacaan, terjemahan,
                mode_pembuatan, contoh_soal, tampilan_bacaan, total_soal
            )
            
            with st.spinner(f"Membuat {total_soal} soal..."):
                if ai_provider == "DeepSeek":
                    hasil, error = buat_soal(api_key, prompt, ai_provider, deepseek_model, "", "", "")
                elif ai_provider == "OpenAI GPT":
                    hasil, error = buat_soal(api_key, prompt, ai_provider, "", openai_model, "", "")
                elif ai_provider == "Maia Router":
                    hasil, error = buat_soal(api_key, prompt, ai_provider, "", "", maia_endpoint, maia_model)
                else:
                    hasil, error = buat_soal(api_key, prompt, ai_provider, "", "", "", "")
            
            if error:
                st.error(f"Gagal: {error}")
            else:
                st.session_state.hasil_soal = hasil
                save_hasil_soal(st.session_state.selected_mapel_id, kelas, topik, hasil)
                st.success(f"Berhasil membuat {total_soal} soal!")
                
                sudah_tampil = False
                for i, soal in enumerate(hasil.get('soal', []), 1):
                    sudah_tampil = display_question(soal, i, tampilan_bacaan, teks_bacaan, terjemahan, sudah_tampil)
                
                if enable_edit and st.session_state.hasil_soal:
                    st.markdown("---")
                    st.subheader("📥 Download")
                    col1, col2 = st.columns(2)
                    with col1:
                        json_str = json.dumps(st.session_state.hasil_soal, indent=2, ensure_ascii=False)
                        st.download_button("📥 Download JSON", json_str, f"soal_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
                        
                        doc = export_to_word(st.session_state.hasil_soal, ki, kd, kelas, topik, sumber_materi, teks_bacaan)
                        doc_bytes = BytesIO()
                        doc.save(doc_bytes)
                        doc_bytes.seek(0)
                        st.download_button("📄 Download Word", doc_bytes, f"soal_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
                    with col2:
                        kunci_doc = buat_kunci_jawaban_word(st.session_state.hasil_soal)
                        kunci_bytes = BytesIO()
                        kunci_doc.save(kunci_bytes)
                        kunci_bytes.seek(0)
                        st.download_button("🔑 Download Kunci", kunci_bytes, f"kunci_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

# ========== MAIN ==========
def main():
    init_session_state()
    
    if not st.session_state.logged_in:
        show_login_page()
    elif st.session_state.selected_mapel_id is None:
        show_dashboard()
    elif not st.session_state.show_generator:
        show_manage_ki_kd()
    else:
        show_generator()

if __name__ == "__main__":
    main()
